"""DOCX text extraction with heading-aware structure."""
from __future__ import annotations

from io import BytesIO
from typing import List, Tuple

from docx import Document


class DocxParser:
    """Extracts structured text from .docx files, preserving paragraph-level
    style information and heading hierarchy."""

    def parse(self, content: bytes) -> List[Tuple[str, str, bool]]:
        """Parse DOCX content and return (text, style_name, is_heading) tuples.

        Each tuple represents one logical paragraph.
        """
        doc = Document(BytesIO(content))
        paragraphs: List[Tuple[str, str, bool]] = []

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            style_name = (para.style.name if para.style else "Normal")
            is_heading = self._is_heading_style(style_name)
            paragraphs.append((text, style_name, is_heading))

        return paragraphs

    def parse_flat(self, content: bytes) -> List[Tuple[int, str]]:
        """Parse DOCX into pages (simulated — DOCX has no native pages).

        Groups paragraphs into simulated pages of ~3000 chars each, similar
        to how they'd render.
        """
        paras = self.parse(content)
        pages: List[Tuple[int, str]] = []
        page_texts: List[str] = []
        current_len = 0
        page_num = 1

        for text, _style, _is_heading in paras:
            if current_len + len(text) > 3000 and page_texts:
                pages.append((page_num, "\n\n".join(page_texts)))
                page_num += 1
                page_texts = []
                current_len = 0
            page_texts.append(text)
            current_len += len(text)

        if page_texts:
            pages.append((page_num, "\n\n".join(page_texts)))

        return pages

    @staticmethod
    def _is_heading_style(style_name: str) -> bool:
        """Check if a Word style name indicates a heading."""
        heading_indicators = [
            "heading", "Heading",
            "title", "Title",
            "subtitle", "Subtitle",
        ]
        for indicator in heading_indicators:
            if indicator in style_name:
                return True

        # Also catch custom heading styles like "H1", "H2"
        if style_name.startswith("H") and len(style_name) <= 3:
            try:
                int(style_name[1:])
                return True
            except ValueError:
                pass

        return False

    def extract_headings(self, content: bytes) -> List[Tuple[str, int]]:
        """Extract headings with their paragraph index from DOCX content."""
        doc = Document(BytesIO(content))
        headings: List[Tuple[str, int]] = []

        for idx, para in enumerate(doc.paragraphs):
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else "Normal"
            if self._is_heading_style(style_name):
                headings.append((text, idx))

        return headings


docx_parser = DocxParser()
